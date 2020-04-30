from __future__ import print_function
import urllib3
from bs4 import BeautifulSoup, SoupStrainer, Comment
import re
import warnings

warnings.filterwarnings('ignore')
DestPathBlog = 'C:\\Users\\wjyzG\\Desktop\\yuuri\\blog\\'

for m in ['201201', '201202', '201203', '201204', '201205', '201206', '201207', '201208', '201209', '201210', '201211',
          '201212',
          '201301', '201302', '201303', '201304', '201305', '201306', '201307', '201308', '201309', '201310', '201311',
          '201312',
          '201401', '201402', '201403', '201404', '201405', '201406', '201407', '201408', '201409', '201410', '201411',
          '201412',
          '201501', '201502', '201503', '201504', '201505', '201506', '201507', '201508', '201509', '201510', '201511',
          '201512',
          '201601', '201602', '201603', '201604', '201605', '201606', '201607', '201608', '201609', '201610', '201611',
          '201612',
          '201701', '201702', '201703', '201704', '201705', '201706', '201707', '201708', '201709', '201710', '201711',
          '201712',
          '201801', '201802', '201803', '201804', '201805', '201806', '201807', '201808', '201809', '201810', '201811',
          '201812',
          '201901', '201902', '201903', '201904', '201905', '201906', '201111', '201112']:
    url = 'http://blog.nogizaka46.com/yuuri.saito/?d=' + str(m)
    http = urllib3.PoolManager()
    response = http.request('GET', url, timeout=5)
    print(m, "response 状态：", response.status)

    soup = BeautifulSoup(response.data, "lxml")
    id_check = soup.find(id="sheet")

    # 爬取博客页数

    page_check = soup.find(class_='paginate')
    page_all = re.findall('\xa0(.*?)\xa0', str(page_check))
    page = len(page_all)
    if page == 0:
        page = 1

    print(m,'共有',page,'页')


    k = 1
    while k <= page:
        url = 'http://blog.nogizaka46.com/yuuri.saito/?p='+str(k)+'&d='+str(m)
        response = http.request('GET', url, timeout=5)
        soup = BeautifulSoup(response.data, "lxml")
        id_check = soup.find(id="sheet")

        # 爬取博客标题

        title_container = id_check.find_all(class_="entrytitle")

        title = []
        for container in title_container:
            tmp = container.text
            title.append(tmp)
        print(m,'_',k,"title_finish")

        # 爬取博客时间

        month_check = soup.find_all(class_='yearmonth')
        date_check = soup.find_all(class_='dd1')
        month = []
        date = []
        for i in month_check:
            temp = i.text
            month.append(temp)
        for i in date_check:
            temp = i.text
            date.append(temp)

        riqi = []
        i = 0
        while i < len(date):
            tmp = month[i] + '/' + date[i]
            tmp = tmp.replace('/', '_')
            riqi.append(tmp)
            i += 1
        print(m , '_' , k , "日期_finish")

        # 爬取文章内容

        body_container = id_check.find_all(class_="entrybody")

        body = []
        for container in body_container:
            tmp = container.text
            body.append(tmp)
        print(m ,'_' ,k , "body_finish")


        # 修改格式输出到word

        i = 0
        while i < len(body):
            body[i] = body[i].replace('\u3000', '\n')
            i += 1

        from docx import Document

        i = 0
        while i < len(body):
            document = Document()
            document.add_heading(title[i], level=1)
            document.add_paragraph(body[i])
            document.save(DestPathBlog + str(m) + '\\' + riqi[i] + '.docx')
            i += 1
        print(m , '_' , k , "word_finish")

        k += 1

    print(m, "finish")
print('all finish')
