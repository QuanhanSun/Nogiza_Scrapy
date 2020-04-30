import os
import pandas as pd
from docx import Document
from docx.shared import Inches
import re

DestPathPic = 'C:\\Users\\wjyzG\\Desktop\\yuuri\\pic\\'
DestPathBlog = 'C:\\Users\\wjyzG\\Desktop\\yuuri\\blog\\'
DestPathBoth = 'C:\\Users\\wjyzG\\Desktop\\yuuri\\both\\'

for m in ['201201', '201202', '201203', '201204', '201205', '201206', '201207', '201208', '201209', '201210', '201211',
          '201212',
          '201301', '201302', '201303', '201304', '201305', '201306', '201307', '201308', '201309', '201310', '201311',
          '201312',
          '201401', '201402', '201403', '201404', '201405', '201406', '201407', '201408', '201409', '201410', '201411',
          '201412',
          '201501', '201502', '201503', '201504', '201505', '201506', '201507', '201508', '201509', '201510', '201511',
          '201512',
          '201601', '201602', '201603', '201604', '201605', '201606', '201607', '201608', '201609', '201610', '201611',
          '201612',
          '201701', '201702', '201703', '201704', '201705', '201706', '201707', '201708', '201709', '201710', '201711',
          '201712',
          '201801', '201802', '201803', '201804', '201805', '201806', '201807', '201808', '201809', '201810', '201811',
          '201812',
          '201901', '201902', '201903', '201904', '201905', '201906', '201111', '201112']:

    pic_path = DestPathPic + str(m) + '\\'

    pic = os.listdir(pic_path)
    pic_date = []

    # 从图片文件名提取出日期

    for date in pic:
        tmp = date[0:10]
        pic_date.append(tmp)

    pic_info = pd.DataFrame({'name': pic, 'date': pic_date})

    # 读入word

    blog_path = DestPathBlog + str(m) + '\\'

    blog = os.listdir(blog_path)
    blog_date = []

    for date in blog:
        tmp = date[0:10]
        blog_date.append(tmp)

    blog_info = pd.DataFrame({'name': blog, 'date': blog_date})

    i = 0
    while i < len(blog):
        j = 0
        trigger = 0
        document = Document(blog_path + blog[i])
        while j < len(pic):
            if pic_info['date'][j] == blog_info['date'][i]:
                trigger = 1
                try:
                    document.add_picture(pic_path + pic[j], width=Inches(3))
                except Exception:
                    pass
                j += 1
            else:
                j += 1
        document.save(DestPathBoth + str(m) + '\\' + blog[i])

        if trigger == 0:
            document = Document(blog_path + blog[i])
            document.save(DestPathBoth + str(m) + '\\' + blog[i])
        print(blog_info['date'][i], 'finish')
        i += 1
